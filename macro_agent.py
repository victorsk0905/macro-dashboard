#!/usr/bin/env python3
"""
================================================================================
 MACRO DASHBOARD AGENT
================================================================================
Pulls the most important global macroeconomic indicators from FRED, builds:
   1. an Excel workbook  (macro_dashboard_YYYY-MM-DD.xlsx)  - the numbers
   2. a chart image      (macro_charts_YYYY-MM-DD.png)       - the trends
   3. a Word summary      (macro_summary_YYYY-MM-DD.docx)     - plain English

USAGE
-----
   export FRED_API_KEY="your_key_here"      # free: https://fredaccount.stlouisfed.org/apikeys
   python macro_agent.py

The agent is self-contained and stateless: every run is a fresh snapshot,
so it is safe to schedule (cron / Task Scheduler / GitHub Actions).
================================================================================
"""

import os
import sys
import datetime as dt
import pandas as pd

from indicators import INDICATORS, GROUP_ORDER

OUTDIR = os.environ.get("MACRO_OUTDIR", os.path.dirname(os.path.abspath(__file__)))
TODAY = dt.date.today().isoformat()


# ------------------------------------------------------------------ DATA LAYER
def get_fred_client():
    key = os.environ.get("FRED_API_KEY", "").strip()
    if not key:
        sys.exit("ERROR: set the FRED_API_KEY environment variable "
                 "(free key at https://fredaccount.stlouisfed.org/apikeys)")
    from fredapi import Fred
    return Fred(api_key=key)


def fetch_series(fred, ind):
    """Return a dict with latest value, prior value, change, and a recent window."""
    raw = fred.get_series(ind["id"]).dropna()
    if raw.empty:
        return None

    # year-over-year transform for price indices
    if ind["transform"] == "yoy_pct":
        # infer periods/year from index frequency
        per_year = _periods_per_year(raw)
        series = (raw.pct_change(per_year) * 100).dropna()
    elif ind["transform"] == "qoq_annualized":
        series = (((raw.pct_change() + 1) ** 4 - 1) * 100).dropna()
    else:
        series = raw

    if series.empty:
        return None

    latest = series.iloc[-1]
    latest_date = series.index[-1].date()
    prior = series.iloc[-2] if len(series) > 1 else float("nan")
    # value 1 year ago (for context on slow series)
    try:
        year_ago = series.asof(series.index[-1] - pd.DateOffset(years=1))
    except Exception:
        year_ago = float("nan")

    return dict(
        ind=ind,
        latest=float(latest),
        latest_date=latest_date,
        prior=float(prior),
        chg=float(latest - prior),
        year_ago=float(year_ago) if pd.notna(year_ago) else float("nan"),
        window=series.tail(60),       # for charting
    )


def _periods_per_year(s):
    if len(s) < 3:
        return 12
    days = (s.index[-1] - s.index[-2]).days
    if days <= 2:   return 252   # daily
    if days <= 10:  return 52    # weekly
    if days <= 45:  return 12    # monthly
    return 4                     # quarterly


# ------------------------------------------------------------- FORMAT HELPERS
def fmt(val, unit):
    if pd.isna(val):
        return "n/a"
    if unit in ("pct",):                 return f"{val:.2f}%"
    if unit in ("index_yoy",):           return f"{val:+.2f}%"
    if unit == "index":                  return f"{val:,.1f}"
    if unit == "count_k":                return f"{val/1000:,.0f}k"
    if unit == "usd":                    return f"${val:,.2f}"
    if unit == "fx":                     return f"{val:,.3f}"
    if unit == "bn_chg_qoq":             return f"{val:,.0f}"
    if unit == "qoq_annualized":         return f"{val:+.2f}%"
    return f"{val:,.2f}"


def direction_word(d):
    if pd.isna(d):            return "flat"
    if abs(d) < 1e-9:         return "flat"
    return "up" if d > 0 else "down"


def health_flag(rec):
    """Green / Amber / Red-ish read for the 'good' direction, else neutral."""
    good = rec["ind"]["good"]
    if good is None:
        return "Context"
    moving = direction_word(rec["chg"])
    if moving == "flat":
        return "Stable"
    return "Improving" if moving == good else "Worsening"


# ----------------------------------------------------------------- SPREADSHEET
def build_workbook(records):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard"

    hdr_fill = PatternFill("solid", fgColor="1F3864")
    grp_fill = PatternFill("solid", fgColor="D9E1F2")
    hdr_font = Font(color="FFFFFF", bold=True, size=11)
    grp_font = Font(bold=True, size=11, color="1F3864")
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    flag_fill = {
        "Improving": PatternFill("solid", fgColor="C6EFCE"),
        "Worsening": PatternFill("solid", fgColor="FFC7CE"),
        "Stable":    PatternFill("solid", fgColor="FFEB9C"),
        "Context":   PatternFill("solid", fgColor="EDEDED"),
    }

    headers = ["Indicator", "Latest", "As of", "Prior", "Change", "1Y ago", "Signal", "What it means"]
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = hdr_fill; cell.font = hdr_font
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        cell.border = border

    by_group = {g: [] for g in GROUP_ORDER}
    for r in records:
        by_group[r["ind"]["group"]].append(r)

    row = 2
    for g in GROUP_ORDER:
        ws.cell(row=row, column=1, value=g).font = grp_font
        for c in range(1, len(headers) + 1):
            ws.cell(row=row, column=c).fill = grp_fill
            ws.cell(row=row, column=c).border = border
        row += 1
        for rec in by_group[g]:
            ind = rec["ind"]
            flag = health_flag(rec)
            vals = [
                ind["label"],
                fmt(rec["latest"], ind["unit"]),
                rec["latest_date"].isoformat(),
                fmt(rec["prior"], ind["unit"]),
                fmt(rec["chg"], ind["unit"]) if ind["unit"] != "count_k" else f"{rec['chg']/1000:+.0f}k",
                fmt(rec["year_ago"], ind["unit"]),
                flag,
                ind["note"],
            ]
            for c, v in enumerate(vals, start=1):
                cell = ws.cell(row=row, column=c, value=v)
                cell.border = border
                cell.alignment = Alignment(vertical="center", wrap_text=(c == 8))
                if c == 7:
                    cell.fill = flag_fill.get(flag, flag_fill["Context"])
                    cell.font = Font(bold=True)
            row += 1

    widths = [30, 12, 12, 12, 12, 12, 12, 60]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    # metadata sheet
    meta = wb.create_sheet("About")
    meta["A1"] = "Macro Dashboard — generated automatically"
    meta["A1"].font = Font(bold=True, size=13)
    notes = [
        f"Generated: {TODAY}",
        "Source: Federal Reserve Economic Data (FRED), St. Louis Fed.",
        "",
        "Signal column meaning:",
        "  Improving  = latest reading moved in the healthy direction",
        "  Worsening  = latest reading moved in the unhealthy direction",
        "  Stable     = essentially unchanged",
        "  Context    = neutral indicator (level matters more than direction)",
        "",
        "This is informational only and is not investment advice.",
    ]
    for i, n in enumerate(notes, start=3):
        meta[f"A{i}"] = n
    meta.column_dimensions["A"].width = 70

    path = os.path.join(OUTDIR, f"macro_dashboard_{TODAY}.xlsx")
    wb.save(path)
    return path


# --------------------------------------------------------------------- CHARTS
def build_charts(records):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    recs = [r for r in records if len(r["window"]) > 2]
    n = len(recs)
    cols = 3
    rows = (n + cols - 1) // cols
    fig, axes = plt.subplots(rows, cols, figsize=(cols * 4.2, rows * 2.6))
    axes = axes.flatten()

    for ax, rec in zip(axes, recs):
        w = rec["window"]
        ax.plot(w.index, w.values, color="#1F3864", linewidth=1.6)
        ax.fill_between(w.index, w.values, w.values.min(), color="#1F3864", alpha=0.06)
        ax.set_title(rec["ind"]["label"], fontsize=9, fontweight="bold")
        ax.tick_params(labelsize=6)
        ax.margins(x=0)
        last = rec["latest"]
        ax.annotate(fmt(last, rec["ind"]["unit"]),
                    xy=(w.index[-1], last), fontsize=7, fontweight="bold",
                    color="#C8472B", ha="right", va="bottom")
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)

    for ax in axes[n:]:
        ax.axis("off")

    fig.suptitle(f"Global Macro Dashboard — recent trends — {TODAY}",
                 fontsize=13, fontweight="bold", y=0.997)
    fig.tight_layout(rect=[0, 0, 1, 0.985])
    path = os.path.join(OUTDIR, f"macro_charts_{TODAY}.png")
    fig.savefig(path, dpi=130)
    plt.close(fig)
    return path


# ------------------------------------------------------- PLAIN-ENGLISH SUMMARY
def narrate_group(group, recs):
    """Build 1-3 plain sentences describing a cluster's state."""
    bits = []
    for r in recs:
        ind = r["ind"]
        v = fmt(r["latest"], ind["unit"])
        mv = direction_word(r["chg"])
        bits.append((ind["label"], v, mv, health_flag(r)))

    if group == "Growth":
        lead = "Growth & jobs: "
    elif group == "Inflation & Money":
        lead = "Inflation: "
    elif group == "Rates & Policy":
        lead = "Interest rates & credit: "
    elif group == "Markets & Currency":
        lead = "Markets & the dollar: "
    else:
        lead = "Rest of the world: "

    sentence = "; ".join(
        f"{label} is {v} ({'rising' if mv=='up' else 'falling' if mv=='down' else 'flat'})"
        for (label, v, mv, _flag) in bits
    )
    return lead + sentence + "."


def overall_read(records):
    """A single top-line paragraph summarizing the whole picture."""
    def find(idx):
        for r in records:
            if r["ind"]["id"] == idx:
                return r
        return None

    parts = []
    infl = find("PCEPILFE") or find("CPIAUCSL")
    if infl:
        lvl = infl["latest"]
        if lvl > 3:    parts.append("inflation remains above the Fed's 2% target")
        elif lvl > 2:  parts.append("inflation is moderately above target but easing toward it")
        else:          parts.append("inflation is at or below the Fed's target")

    un = find("UNRATE")
    if un:
        parts.append(f"unemployment is {un['latest']:.1f}% and "
                     f"{'rising' if un['chg']>0 else 'steady-to-falling'}")

    curve = find("T10Y2Y")
    if curve:
        if curve["latest"] < 0:
            parts.append("the yield curve is inverted, a classic late-cycle warning")
        else:
            parts.append("the yield curve is positive (no recession flag from this gauge)")

    hy = find("BAMLH0A0HYM2")
    if hy:
        if hy["latest"] > 5:
            parts.append("credit spreads are elevated, signaling market stress")
        else:
            parts.append("credit spreads are calm, so markets are not pricing distress")

    dxy = find("DTWEXBGS")
    if dxy:
        parts.append(f"the dollar is {'softening' if dxy['chg']<0 else 'firming'}")

    return ("In one line: " + "; ".join(parts) + ". "
            "Read the cluster notes below for the detail.")


def build_word(records, chart_path, analysis=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    VERM = RGBColor(0xC8, 0x47, 0x2B)

    doc = Document()
    style = doc.styles["Normal"].font
    style.name = "Calibri"; style.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("Global Macro Dashboard")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sr = sub.add_run(f"Weekly snapshot · {TODAY} · source: FRED (St. Louis Fed)")
    sr.italic = True; sr.font.size = Pt(10); sr.font.color.rgb = VERM

    # ---- Claude's weekly narrative (if available) sits at the very top
    if analysis:
        h = doc.add_paragraph(); hr = h.add_run("Weekly read")
        hr.bold = True; hr.font.size = Pt(14); hr.font.color.rgb = NAVY

        hl = doc.add_paragraph()
        hlr = hl.add_run(analysis["headline"])
        hlr.bold = True; hlr.font.size = Pt(12)

        reg = doc.add_paragraph()
        rl = reg.add_run("Regime: "); rl.bold = True
        reg.add_run(analysis["regime"])

        for para in str(analysis["narrative"]).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        wh = doc.add_paragraph(); whr = wh.add_run("What to watch next")
        whr.bold = True; whr.font.size = Pt(12); whr.font.color.rgb = VERM
        for item in analysis["watch"]:
            doc.add_paragraph(str(item), style="List Bullet")

        cred = doc.add_paragraph()
        cr = cred.add_run("Narrative generated by Claude from this week's data.")
        cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(0x80,0x80,0x80)

    # ---- rule-based top-line read (always present; primary if no Claude layer)
    h = doc.add_paragraph()
    hr = h.add_run("Data summary" if analysis else "The big picture")
    hr.bold = True; hr.font.size = Pt(14); hr.font.color.rgb = NAVY
    doc.add_paragraph(overall_read(records))

    # ---- chart
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- cluster-by-cluster plain English
    by_group = {g: [] for g in GROUP_ORDER}
    for r in records:
        by_group[r["ind"]["group"]].append(r)

    h = doc.add_paragraph(); hr = h.add_run("What each part of the economy is saying")
    hr.bold = True; hr.font.size = Pt(14); hr.font.color.rgb = NAVY

    for g in GROUP_ORDER:
        recs = by_group[g]
        if not recs:
            continue
        gh = doc.add_paragraph(); ghr = gh.add_run(g)
        ghr.bold = True; ghr.font.size = Pt(12); ghr.font.color.rgb = VERM
        doc.add_paragraph(narrate_group(g, recs))
        # a compact "movers" line
        worsening = [r["ind"]["label"] for r in recs if health_flag(r) == "Worsening"]
        improving = [r["ind"]["label"] for r in recs if health_flag(r) == "Improving"]
        mv = doc.add_paragraph()
        if improving:
            a = mv.add_run("Improving: "); a.bold = True
            mv.add_run(", ".join(improving) + ".   ")
        if worsening:
            b = mv.add_run("Worsening: "); b.bold = True
            mv.add_run(", ".join(worsening) + ".")
        if not improving and not worsening:
            mv.add_run("No notable directional moves this period.").italic = True

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("This document is generated automatically for information only. "
                      "It is not investment advice.")
    fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x80,0x80,0x80)

    path = os.path.join(OUTDIR, f"macro_summary_{TODAY}.docx")
    doc.save(path)
    return path


# ------------------------------------------------------------- HTML DASHBOARD
def build_dashboard(records):
    """Render the interactive single-file HTML dashboard from a template."""
    import json
    tpl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "dashboard_template.html")
    if not os.path.exists(tpl_path):
        print("  (dashboard_template.html not found — skipping HTML dashboard)")
        return None

    payload = []
    for r in records:
        ind = r["ind"]
        payload.append(dict(
            id=ind["id"], label=ind["label"], group=ind["group"], unit=ind["unit"],
            note=ind["note"], good=ind["good"],
            latest=round(r["latest"], 3), prior=round(r["prior"], 3),
            chg=round(r["chg"], 3),
            latest_date=r["latest_date"].isoformat(),
            latest_fmt=fmt(r["latest"], ind["unit"]),
            prior_fmt=fmt(r["prior"], ind["unit"]),
            flag=health_flag(r),
            spark=[round(float(x), 3) for x in r["window"].tail(40).tolist()],
        ))

    tpl = open(tpl_path).read()
    html = tpl.replace("/*__DATA__*/[]", json.dumps(payload))
    path = os.path.join(OUTDIR, f"macro_dashboard_{TODAY}.html")
    with open(path, "w") as f:
        f.write(html)
    return path


# ------------------------------------------------------------------------ MAIN
def main():
    fred = get_fred_client()
    print("Fetching indicators from FRED ...")
    records = []
    for ind in INDICATORS:
        try:
            rec = fetch_series(fred, ind)
            if rec:
                records.append(rec)
                print(f"  ok  {ind['label']:<34} {fmt(rec['latest'], ind['unit'])}")
            else:
                print(f"  --  {ind['label']:<34} no data")
        except Exception as e:
            print(f"  ERR {ind['label']:<34} {e}")

    if not records:
        sys.exit("No data fetched — check API key / network.")

    xlsx = build_workbook(records)
    png  = build_charts(records)

    print("Generating Claude weekly narrative ...")
    from claude_analysis import generate_analysis
    analysis = generate_analysis(records, fmt, TODAY)

    docx = build_word(records, png, analysis)
    html = build_dashboard(records)

    print("\nDone. Files written:")
    for p in (html, docx, xlsx, png):
        print("  ", p)


if __name__ == "__main__":
    main()
